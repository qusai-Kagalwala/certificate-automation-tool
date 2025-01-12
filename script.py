import pandas as pd
from docx import Document
from docxcompose.composer import Composer

# Function to fill in the placeholders in the Word template
def fill_certificate(template_path, output_path, data):
    doc = Document(template_path)

    # Loop through paragraphs in the document
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            # Replace placeholders with the actual data
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    # Save the modified document to the output path
    doc.save(output_path)

# Function to generate certificates from the CSV
def generate_certificates_from_csv(csv_path, template_path):
    # Read the CSV file into a DataFrame
    df = pd.read_csv(csv_path)
    generated_files = []

    # Loop through each row in the DataFrame
    for idx, row in df.iterrows():
        # Prepare the data for placeholder replacement
        data = {
            '[Name]': row['Name'],
            '[College]': row['College'],
            '[Event]': row['Event'],
        }

        # Define the output file name for each certificate
        output_path = f'output/certificate_{idx + 1}.docx'
        generated_files.append(output_path)

        # Fill the certificate with data and save it
        fill_certificate(template_path, output_path, data)
    
    return generated_files

# Function to combine all Word files into a single document
def combine_word_files(output_files, combined_file_path):
    # Create a base document
    combined_document = Document(output_files[0])
    composer = Composer(combined_document)

    # Append other documents to the base document
    for file_path in output_files[1:]:
        doc = Document(file_path)
        composer.append(doc)

    # Save the combined document
    composer.save(combined_file_path)

# Entry point of the script
if __name__ == '__main__':
    import os

    # Define file paths
    csv_path = 'info.csv'  # Path to the CSV file with participant info
    template_path = 'template.docx'  # Path to the Word template
    combined_file_path = 'combined_certificates.docx'  # Output for combined Word file

    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)

    # Generate certificates from the CSV using the provided template
    generated_files = generate_certificates_from_csv(csv_path, template_path)

    # Combine all generated Word files into a single document
    combine_word_files(generated_files, combined_file_path)

    print(f"Combined Word file saved as '{combined_file_path}'")

